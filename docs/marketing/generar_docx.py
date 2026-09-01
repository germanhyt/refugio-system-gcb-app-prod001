"""Genera el informe SEO/GEO en .docx a partir del markdown."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD = Path(__file__).parent / "INFORME-SEO-GEO.md"
OUT = Path(__file__).parent / "INFORME-SEO-GEO.docx"

ACCENT = RGBColor(0xB0, 0x7A, 0x1E)      # color de marca tipo mostaza/ámbar
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)


def add_runs(paragraph, text):
    """Agrega texto con **negritas** y `código` inline."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                r = paragraph.add_run(m.group(1))
                r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
                r.underline = True
        else:
            paragraph.add_run(part)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def main():
    doc = Document()

    # Estilos base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for level, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11.5, DARK),
    ]:
        st = doc.styles[level]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True

    lines = MD.read_text(encoding="utf-8").splitlines()

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Informe de SEO y GEO")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Refugio Gastronómico — refugiogastronomico.pe")
    r.font.size = Pt(14)
    r.font.color.rgb = GREY
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Equipo de Marketing · Septiembre 2026 · Fase 1 implementada")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    doc.add_paragraph()

    i = 0
    n = len(lines)
    in_meta = False
    while i < n:
        line = lines[i]

        # Saltar el bloque de metadatos inicial del markdown
        if line.strip().startswith("**Proyecto:**"):
            in_meta = True
            i += 1
            continue
        if in_meta:
            if line.strip() == "---":
                in_meta = False
            i += 1
            continue

        # Regla horizontal
        if line.strip() == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "B07A1E")
            pbdr.append(bottom)
            pPr.append(pbdr)
            i += 1
            continue

        # Encabezados
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # # 1 -> Heading 1 (en el doc, H1 del md ya se usó como portada)
            if level == 1:
                doc.add_heading(text, level=1)
            else:
                doc.add_heading(text, level=level)
            i += 1
            continue

        # Tablas
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for j, h in enumerate(header):
                hdr[j].text = ""
                run = hdr[j].paragraphs[0].add_run(re.sub(r"\*\*|\*", "", h))
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(hdr[j], "B07A1E")
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    if j < len(cells):
                        cells[j].text = ""
                        add_runs(cells[j].paragraphs[0], re.sub(r"\*\*", "", val))
                        for r in cells[j].paragraphs[0].runs:
                            r.font.size = Pt(9.5)
            doc.add_paragraph()
            continue

        # Listas
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text)
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_runs(p, text)
            i += 1
            continue

        # Bloque de código
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "F4F1EA")
            pPr.append(shd)
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Párrafo normal
        if line.strip() == "":
            i += 1
            continue
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    # Pie de página
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Refugio Gastronómico · Informe SEO/GEO · Fase 1")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
